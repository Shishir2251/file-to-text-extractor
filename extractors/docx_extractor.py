"""
DOCX Extractor
Extract text from Word documents
"""

import docx
import io
import logging

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """Extract text from DOCX files"""
    
    def extract(self, file_bytes: bytes) -> str:
        """
        Extract text from DOCX
        
        Args:
            file_bytes: DOCX file as bytes
        
        Returns:
            Extracted text
        """
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            
            # Extract text from paragraphs
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
            
            # Combine all text
            text = "\n".join(paragraphs)
            if table_text:
                text += "\n\n" + "\n".join(table_text)
            
            if not text.strip():
                raise Exception("No text found in DOCX file")
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise Exception(f"DOCX extraction failed: {str(e)}")